"""Ingest adapter tests: local notes, Notion zip, Drive docx/gdoc, and the
end-to-end ingest script against a synthetic corpus."""

from __future__ import annotations

import zipfile
from pathlib import Path

import pytest

from adonis.ingest.gdrive import parse_docx, parse_gdoc
from adonis.ingest.local_notes import parse_local_note
from adonis.ingest.notion import parse_notion_export


def _make_notion_zip(path: Path) -> None:
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("Home.md", "# Home\n\nWelcome **home**.\n\n![banner](img.png)\n")
        zf.writestr("Projects/Atlas.md", "# Atlas\n\n## Status\n\nAtlas ships in **March**.\n")
        zf.writestr("Notion Databases/Tasks.csv", "id,status,name\n1,open,Onboard\n2,done,Launch\n")
        zf.writestr("img.png", b"fake-binary")
    assert path.is_file()


def _make_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Meeting minutes draft.")
    doc.add_paragraph("Decision: use Postgres.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "key"
    table.rows[0].cells[1].text = "value"
    core = doc.core_properties
    core.title = "Minutes 2026-01-05"
    doc.save(str(path))


# --- local notes ------------------------------------------------------------


def test_parse_local_markdown(tmp_path):
    p = tmp_path / "note.md"
    p.write_text("# Hi\n\n**bold** text.\n", encoding="utf-8")
    rec = parse_local_note(p)
    assert rec.source == "local"
    assert rec.title == "note"
    assert rec.format == "md"
    assert rec.raw_text == "Hi\n\nbold text."
    assert "**" not in rec.raw_text


def test_parse_local_txt_no_strip(tmp_path):
    p = tmp_path / "plain.txt"
    p.write_text("just plain **text**", encoding="utf-8")
    rec = parse_local_note(p)
    assert rec.raw_text == "just plain **text**"
    assert rec.format == "txt"


# --- notion zip -------------------------------------------------------------


def test_parse_notion_zip(tmp_path):
    zip_path = tmp_path / "Notion_Export.zip"
    _make_notion_zip(zip_path)
    records = parse_notion_export(zip_path)
    assert len(records) == 3
    by_title = {r.title: r for r in records}
    assert "Home" in by_title and "Atlas" in by_title and "Tasks" in by_title
    atlas = by_title["Atlas"]
    assert atlas.source == "notion"
    assert "**March**" not in atlas.raw_text and "March" in atlas.raw_text
    tasks = by_title["Tasks"]
    assert tasks.format == "csv"
    assert tasks.metadata["columns"] == ["id", "status", "name"]
    assert tasks.metadata["row_count"] == 2


# --- gdrive -----------------------------------------------------------------


def test_parse_docx(tmp_path):
    p = tmp_path / "minutes.docx"
    _make_docx(p)
    rec = parse_docx(p)
    assert rec.source == "gdrive"
    assert rec.title == "Minutes 2026-01-05"
    assert "Decision: use Postgres." in rec.raw_text
    assert "key | value" in rec.raw_text


def test_parse_gdoc_returns_none(tmp_path):
    p = tmp_path / "doc.gdoc"
    p.write_text('{"id": "abc123", "url": "https://docs.google.com/document/d/abc123/edit"}', encoding="utf-8")
    assert parse_gdoc(p) is None


def test_parse_gdoc_invalid_json_raises(tmp_path):
    p = tmp_path / "broken.gdoc"
    p.write_text("not json", encoding="utf-8")
    with pytest.raises(ValueError):
        parse_gdoc(p)


# --- end-to-end script ------------------------------------------------------


def test_ingest_script_end_to_end(tmp_path, monkeypatch):
    from adonis import config as cfg

    monkeypatch.setenv("ADONIS_DB_PATH", str(tmp_path / "db" / "adonis.sqlite"))
    monkeypatch.setenv("ADONIS_CORPUS_DIR", str(tmp_path / "corpus"))
    monkeypatch.setenv("ADONIS_REPORTS_DIR", str(tmp_path / "reports"))
    monkeypatch.setenv("ADONIS_LLM_API_KEY", "test-key")
    cfg._settings = None

    corpus = tmp_path / "corpus"
    corpus.mkdir()
    (corpus / "notes").mkdir()
    (corpus / "notes" / "alpha.md").write_text("# Alpha note\n\nsome **content**\n", encoding="utf-8")
    (corpus / "notes" / "beta.txt").write_text("plain text file", encoding="utf-8")
    _make_notion_zip(corpus / "export.zip")
    _make_docx(corpus / "minutes.docx")
    (corpus / "drive").mkdir()
    (corpus / "drive" / "doc.gdoc").write_text('{"id": "abc"}', encoding="utf-8")
    (corpus / "readme.pdf").write_bytes(b"%PDF fake")

    from adonis.db import get_conn
    from scripts.ingest_corpus import run

    stats = run(corpus)
    assert stats.files_seen == 6
    assert stats.inserted == 6  # 2 local + 3 notion + 1 docx
    assert stats.failed == 0
    assert stats.skipped == 1  # gdoc
    assert stats.ignored == 1  # pdf
    assert stats.doc_count == 6

    conn = get_conn()
    try:
        rows = conn.execute("SELECT source, format FROM documents").fetchall()
        sources = sorted((r["source"], r["format"]) for r in rows)
        assert ("gdrive", "docx") in sources
        assert ("local", "md") in sources
        assert ("local", "txt") in sources
        assert ("notion", "csv") in sources
        assert ("notion", "md") in sources
    finally:
        conn.close()

    # Re-run: everything is a duplicate, nothing is re-inserted.
    stats2 = run(corpus)
    assert stats2.inserted == 0
    assert stats2.duplicates == 6
    assert stats2.failed == 0
    cfg._settings = None
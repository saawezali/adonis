"""Google Drive export parser (.docx, .gdoc).

- .docx: real content, parsed with python-docx (paragraphs + tables).
- .gdoc: a JSON shortcut file that only contains the doc id and URL; the
  content is not available offline. Returns None so the caller can report it
  as skipped rather than storing an empty document.
"""

from __future__ import annotations

import json
from pathlib import Path

from adonis.ingest.base import DocumentRecord
from adonis.normalize.text import content_hash, normalize_text


def parse_docx(path: Path) -> DocumentRecord:
    """Parse a .docx file into a DocumentRecord.

    Body paragraphs are joined in order; tables are rendered as pipe-joined
    rows so fact-finding survives table layouts.
    """
    from docx import Document  # lazy import: keeps ingest CLI startup fast

    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    text = normalize_text("\n".join(p for p in parts if p.strip()))
    core = doc.core_properties
    return DocumentRecord(
        source="gdrive",
        source_id=content_hash(text),
        title=(core.title or path.stem),
        path=str(path),
        format="docx",
        raw_text=text,
        metadata={"filename": path.name},
    )


def parse_gdoc(path: Path) -> DocumentRecord | None:
    """Parse a .gdoc shortcut file. Returns None (no offline content).

    Raises ValueError if the file looks like a .gdoc but is not valid JSON.
    """
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict) or "id" not in data:
        return None
    return None